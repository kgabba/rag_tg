import os
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import io
from typing import List
from pypdf import PdfReader
import docx  # python-docx

API_KEY = os.getenv("API_KEY")
#BASE_URL = "https://aleron-llm.neuraldeep.tech/"

# LLM (для диалогов, не обязателен для эмбеддингов)
llm = ChatOpenAI(
    model="gpt-4o-mini",
    api_key=API_KEY,
    #base_url=BASE_URL,
    temperature=0
)

# Модель эмбеддингов 
embeddings_model = OpenAIEmbeddings(
    model="text-embedding-3-small",   
    api_key=API_KEY
    #base_url=BASE_URL,
)

# Разбивка текста на чанки
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,     # размер чанка (символы)
    chunk_overlap=200,  # перекрытие
)

def chunk_and_embed_to_db(text: str, conn):
    docs = text_splitter.create_documents([text])
    chunks = [d.page_content for d in docs]

    vectors = embeddings_model.embed_documents(chunks)

    cursor = conn.cursor()
    for chunk, vec in zip(chunks, vectors):
        # pgvector принимает вот такой строчный формат: [1,2,3,...]
        vec_str = "[" + ",".join(str(x) for x in vec) + "]"
        cursor.execute(
            """
            INSERT INTO embeddings (text, embedding)
            VALUES (%s, %s::vector)
            """,
            (chunk, vec_str),
        )

    return len(chunks)

def get_top_k_chunks(question: str, conn, k: int = 3) -> list[str]:
    """Векторизуем вопрос и достаём top-k ближайших чанков из таблицы embeddings."""
    # один вектор для запроса
    q_vec = embeddings_model.embed_query(question)
    q_vec_str = "[" + ",".join(str(x) for x in q_vec) + "]"

    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT text
        FROM embeddings
        ORDER BY embedding <-> %s::vector
        LIMIT %s
        """,
        (q_vec_str, k),
    )
    rows = cursor.fetchall()
    return [r[0] for r in rows]


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Простое извлечение текста из PDF."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text: List[str] = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt:
            pages_text.append(txt)
    return "\n\n".join(pages_text)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Извлекаем параграфы и таблицы из DOCX."""
    document = docx.Document(io.BytesIO(file_bytes))
    parts: List[str] = []

    # Параграфы
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # Таблицы – разворачиваем строки в текст
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Определяем тип файла по расширению и вытаскиваем текст."""
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    if lower_name.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)

    # если захотим – потом добавим .doc, .txt и т.д.
    raise ValueError("Поддерживаются только .pdf и .docx")


def chunk_and_embed_file(file_bytes: bytes, filename: str, conn) -> int:
    """Читает файл → вытаскивает текст → режет на чанки → пишет эмбеддинги в БД."""
    text = extract_text_from_bytes(file_bytes, filename)
    if not text.strip():
        return 0

    return chunk_and_embed_to_db(text, conn)